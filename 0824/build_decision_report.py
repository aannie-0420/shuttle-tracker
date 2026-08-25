from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT=Path(r"D:\shuttle\0824\DOC\Advantech_ShuttleFlow_總務ESG決策分析報表.docx")
BLUE="2E74B5"; DARK="1F4D78"; LIGHT="E8EEF5"; GRAY="F2F4F7"; GREEN="EAF4EA"; GOLD="FFF5D6"; RED="FCE8E6"

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for name,val in [("top",80),("start",120),("bottom",80),("end",120)]:
        n=tcMar.find(qn("w:"+name))
        if n is None: n=OxmlElement("w:"+name); tcMar.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")

def geometry(t,widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    pr=t._tbl.tblPr; tw=pr.find(qn("w:tblW"))
    if tw is None: tw=OxmlElement("w:tblW"); pr.append(tw)
    tw.set(qn("w:w"),str(sum(widths))); tw.set(qn("w:type"),"dxa")
    ind=pr.find(qn("w:tblInd"))
    if ind is None: ind=OxmlElement("w:tblInd"); pr.append(ind)
    ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa")
    grid=t._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for w in widths:
        c=OxmlElement("w:gridCol"); c.set(qn("w:w"),str(w)); grid.append(c)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            tcPr=c._tc.get_or_add_tcPr(); tcw=tcPr.find(qn("w:tcW"))
            if tcw is None: tcw=OxmlElement("w:tcW"); tcPr.append(tcw)
            tcw.set(qn("w:w"),str(widths[i])); tcw.set(qn("w:type"),"dxa"); margins(c)
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc,headers,rows,widths,fill=LIGHT):
    t=doc.add_table(rows=1,cols=len(headers)); geometry(t,widths)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,fill); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
        r=p.add_run(h); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(11,37,69)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0)
            p.add_run(str(v)).font.size=Pt(8.2)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def bullet(doc,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); p.add_run(text)

def numbered(doc,text):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.space_after=Pt(3); p.add_run(text)

def callout(doc,label,text,fill=GRAY):
    t=doc.add_table(rows=1,cols=1); geometry(t,[9360]); c=t.cell(0,0); shade(c,fill)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+"："); r.bold=True; r.font.color.rgb=RGBColor(31,77,120); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def page_field(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=p.add_run("第 ")
    a=OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"),"begin")
    b=OxmlElement("w:instrText"); b.set(qn("xml:space"),"preserve"); b.text=" PAGE "
    c=OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"),"end")
    r._r.append(a); r._r.append(b); r._r.append(c); p.add_run(" 頁")

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
sec.header_distance=Inches(.3); sec.footer_distance=Inches(.3)
styles=doc.styles; n=styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(10.5); n.font.color.rgb=RGBColor(35,45,55); n.paragraph_format.space_after=Pt(5); n.paragraph_format.line_spacing=1.12
for name,size,color,before,after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",11.5,DARK,8,4)]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
h=sec.header.paragraphs[0]; h.text="Advantech ShuttleFlow｜總務／ESG 決策分析報表"; h.runs[0].font.size=Pt(8); h.runs[0].font.color.rgb=RGBColor(100,110,120)
f=sec.footer.paragraphs[0]; f.add_run(); f.runs[0].font.size=Pt(8); f.runs[0].font.color.rgb=RGBColor(100,110,120); page_field(f)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run("Advantech ShuttleFlow"); r.font.size=Pt(25); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14); r=p.add_run("總務／ESG 決策分析報表規格"); r.font.size=Pt(19); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph("目的：將 Firebase 即時車況、ETA、GPS 軌跡與班次資料，轉換成可支援總務營運決策與 ESG 成效說明的管理報表。")
callout(doc,"管理重點","本報表不只回答「現在車在哪裡」，還要回答「服務是否穩定、同仁是否等得更少、車輛是否配置合理、是否有可被量化的減碳效益」。",GREEN)

doc.add_heading("一、報表要支援的決策問題",1)
for x in ["總務：尖峰時段是否有足夠車輛？哪一條路線或哪個時段最容易延誤？","總務：目前的班表、車輛數與車型配置是否需要調整？","ESG：接駁服務是否提高使用率、降低私人載具需求，並能形成可追溯的減碳佐證？","管理階層：系統上線後，等待時間、準點率、服務可用率是否改善？","維運：GPS、網路、Google ETA 或駕駛端是否有異常？"]: bullet(doc,x)

doc.add_heading("二、建議報表架構",1)
add_table(doc,["報表模組","主要回答問題","建議頻率","使用者"],[
["即時營運總覽","現在有哪些車在線、方向與 ETA 是否可信？","即時／每1分鐘","總務、值班"],
["服務品質月報","在線率、GPS 可用率、ETA 準確度如何？","每日／每月","總務主管"],
["車輛與班次分析","哪個時段需要增車或調整班距？","每月／季度","總務、採購"],
["ESG 成效報表","接駁使用量與減碳效益如何估算？","每月／年度","ESG"],
["資料品質報表","哪些資料不完整或影響決策？","每日／每月","IT、維運"]],[1900,3500,1500,2460])

doc.add_heading("三、可直接產出的管理 KPI",1)
add_table(doc,["KPI","計算公式","主要 Firebase 參數","決策用途"],[
["即時在線率","在線車輛數 ÷ 應服務車輛數 ×100%","status、gpsUpdatedAt、heartbeatAt","判斷即時車是否足夠"],
["GPS 可用率","GPS 未超過60秒時間 ÷ 應服務時間 ×100%","gpsUpdatedAt、timestamp、status","判斷即時資訊可靠度"],
["Google ETA 成功率","success 週期 ÷ 總週期 ×100%","googleCycleStatus、googleCycleId、googleError","評估 ETA 服務穩定度"],
["路線偏離率","offRoute 點數 ÷ 有效 GPS 點數 ×100%","routeStatus、lat、lng、accuracy","評估路線與駕駛執行"],
["平均車程","有效行程車程總和 ÷ 有效行程數","archive、stopStatus、timestamp","調整班表與車程設定"],
["車輛利用率","服務時間 ÷ 可用營運時間 ×100%","sessionId、status、archive","評估車輛是否閒置或不足"],
["ETA 誤差","實際抵達時間－預估抵達時間","etaToCampus、etaCycles、actual arrival*","評估 ETA 可信度"]],[1600,2650,2750,2360])
doc.add_paragraph("* ETA 誤差需要額外保存實際抵達事件；目前資料可支援計算基礎，但不一定已完整保存 actual arrival。").italic=True

doc.add_heading("四、Firebase 參數與分析用途",1)
add_table(doc,["Firebase 參數／路徑","可產出的分析","應用或計算方式"],[
["/shuttle/small、/shuttle/medium","即時車況、在線車輛、最快車","status=online 且現在時間-gpsUpdatedAt≤60秒；依 totalCommute 由小到大排序"],
["lat、lng、accuracy","行駛路線、定位品質、偏離路線","以座標重建軌跡；accuracy 過大時標記低可信度"],
["direction、stopStatus","去程／回程、站點狀態","依 direction 顯示前往 A8／返回園區；依 stopStatus 切換 ETA"],
["etaToCampus、totalCommute","到園區 ETA、等車加車程","totalCommute=等待時間+車程；按狀態加總不同路段"],
["currentToA8、currentToCampus、a8ToCampus、campusToA8","分段車程、往返總時間","toA8：currentToA8+a8ToCampus；返回園區：currentToCampus；再加 campusToA8"],
["etaSource、etaStatus、routeStatus","資料可信度標籤","google／coordinate／schedule；offRoute 或過期時降級"],
["googleUpdatedAt、googleGpsUpdatedAt、googleCycleStatus","Google ETA 新鮮度與成功率","Google ETA 5分鐘內且 GPS 對齊90秒內才列為新鮮"],
["sharedTraffic.minutes、updatedAt、status","路況快取與趨勢","未超過10分鐘直接使用；過期才查詢；failed 使用備援值"],
["archive GPS points","歷史車程、停留、偏離、GPS 中斷","依日期、車型、sessionId 分組；首末站點時間計算實際車程"],
["etaCycles/{cycleId}/{type}","ETA 版本追蹤、車輛公平比較","同一 cycleId 的結果才適合直接比較；統計成功、失敗與波動"]],[3000,2600,3760])

doc.add_heading("五、總務應查看的營運分析",1)
doc.add_heading("5.1 即時營運總覽",2)
for x in ["在線車輛數／應服務車輛數。","各車輛方向：前往 A8、返回園區、已抵達、依班表行駛。","每台車的 ETA、總行程時間與 ETA 來源。","目前最快車與被排除原因，例如 GPS 超過60秒、方向不符、路線偏離。","Google ETA、座標估算或班表的資料狀態。"]: bullet(doc,x)
callout(doc,"建議畫面","第一區顯示「現在可搭哪一台車」；第二區顯示「目前資料是否可信」；第三區顯示「異常需要誰處理」。",GOLD)

doc.add_heading("5.2 班次與車輛配置分析",2)
add_table(doc,["分析項目","計算方式","可做的決策"],[
["尖峰需求","依時段統計在線車輛、班次、搭乘事件或查詢量","增加班次或增派車輛"],
["車型配置","比較 small／medium 的服務時間、行程數與利用率","車型調度與租賃配置"],
["班距合理性","實際服務間隔與班表間隔比較","縮短尖峰班距或調整末班車"],
["延誤熱點","依時段、方向、路段比較實際車程與 ETA 誤差","調整發車緩衝或路線"],
["服務可用性","GPS 有效率、在線率、資料中斷時間","改善駕駛裝置、網路與流程"]],[1900,3900,3560])

doc.add_heading("六、ESG 應查看的分析",1)
doc.add_heading("6.1 現有資料可先產出的指標",2)
for x in ["接駁車服務趟次：由 sessionId、type、timestamp 與行程完成狀態統計。","接駁車行駛里程：由 archive GPS 座標逐點累加距離。","車型使用分布：由 type 與車輛能源資料對照。","尖峰服務覆蓋率：以應服務時段中 GPS 有效且在線的時間計算。","空駛或低利用率時段：以行駛軌跡、服務時間與搭乘事件交叉分析。"]: bullet(doc,x)
doc.add_heading("6.2 需補充的 ESG 資料",2)
callout(doc,"重要限制","目前 Firebase 的 GPS 與車況資料可以支援行駛量與服務量；若要正式宣告減碳量，仍需補充搭乘人數、車輛能源／油耗、私人載具替代基準等資料。",RED)
add_table(doc,["資料","用途","建議新增欄位"],[
["搭乘人數","估算被替代私人載具人次","passengerCount、boardingEvent"],
["車輛能源與油耗","計算接駁車本身排放","fuelType、fuelUsed、fuelEfficiency"],
["實際抵達事件","計算 ETA 準確度與準點率","actualArrivedAtA8、actualArrivedAtCampus"],
["行程完成事件","確認班次是否完成","tripStartedAt、tripCompletedAt、tripStatus"],
["查詢／搭乘事件","估算等待時間與使用需求","employeeSessionId、selectedVehicle、boardedAt"],
["私人載具替代基準","計算減碳基準線","baselineDistance、emissionFactor"]],[2200,3550,3610])

doc.add_heading("七、建議的減碳計算方式",1)
doc.add_paragraph("建議將接駁車排放與被替代的私人載具排放分開計算，避免只用接駁車行駛里程推估減碳。")
for x in ["被替代私人載具排放 = 搭乘人數 × 每人替代距離 × 私人載具排放係數。","接駁車排放 = 實際行駛距離 × 車輛排放係數；若有油耗，則用耗油量 × 燃料排放係數。","估算減碳量 = 被替代私人載具排放 - 接駁車排放。","減碳率 = 估算減碳量 ÷ 被替代私人載具排放 ×100%。"]: numbered(doc,x)
doc.add_paragraph("示例：20 人各替代10公里私人開車，私人載具係數0.18 kg CO₂e/km，私人載具排放為36 kg CO₂e。若接駁車本趟排放15 kg CO₂e，估算減碳量為21 kg CO₂e。")

doc.add_heading("八、報表資料品質與治理規則",1)
add_table(doc,["檢核項目","建議規則","報表顯示"],[
["GPS 新鮮度","gpsUpdatedAt 超過60秒不列入即時推薦","離線／資料過期"],
["Google ETA 新鮮度","googleUpdatedAt 超過5分鐘或 GPS 對齊超過90秒","改用座標估算"],
["定位精度","accuracy 過大時不作為高可信度軌跡","定位精度不足"],
["路線狀態","routeStatus=offRoute 時不直接採用 Google ETA","路線重新估算"],
["資料缺口","沒有 actual arrival、搭乘人數時，不宣稱 ETA 準確度或正式減碳量","資料不足，僅供趨勢"],
["隱私","lat/lng、deviceId、sessionId 只供授權人員查詢","管理報表採彙總"]],[1900,4700,2760])

doc.add_heading("九、建議月報版型",1)
add_table(doc,["頁面","內容","圖表建議"],[
["第1頁：管理摘要","服務趟次、在線率、平均車程、ETA 準確度、估算減碳量","KPI 卡片＋本月結論"],
["第2頁：服務品質","按日／時段的在線率、GPS 可用率、Google ETA 成功率","折線圖＋異常日"],
["第3頁：車輛與班次","小車／中車行程、利用率、尖峰時段與方向需求","柱狀圖＋熱點矩陣"],
["第4頁：路線與車程","A8→園區、園區→A8 的平均／P90 車程與延誤","趨勢圖或箱型圖"],
["第5頁：ESG","搭乘人次、行駛公里、替代私人載具、減碳量","瀑布圖或年度累計"],
["第6頁：資料品質","過期 GPS、offRoute、API 失敗、資料缺口","異常清單＋責任人"]],[1600,4700,3060])

doc.add_heading("十、執行優先順序",1)
for x in ["第一階段：用現有 Firebase 欄位建立即時營運總覽、在線率、GPS 可用率、車程與路線分析。","第二階段：補存 actual arrival、trip completion、passengerCount，建立 ETA 準確度與搭乘需求分析。","第三階段：補上車輛能源／油耗與私人載具替代基準，形成可稽核的 ESG 減碳報表。","第四階段：建立月報自動化，提供總務、ESG 與管理階層不同權限與彙總層級。"]: numbered(doc,x)
callout(doc,"結論","Firebase 現有資料已足以支援即時車況、服務可用性、車程與車輛配置的管理決策；若要支援正式 ESG 減碳與員工體驗成效，最優先應補上實際抵達事件、搭乘人數、班次完成事件、車輛能源／油耗四類資料。",GREEN)

doc.add_heading("附錄：本報表依據的程式與資料範圍",1)
for x in [r"D:\shuttle\0824\driver.html：駕駛端即時 GPS、ETA、路況與歷史資料寫入。",r"D:\shuttle\0824\shuttle-driver-runtime.js：駕駛端 owner、heartbeat、GPS archive 與批次上傳。",r"D:\shuttle\0824\shuttle-index-runtime.js：員工端在線判斷、ETA 來源選擇、最快車與班表備援。",r"D:\shuttle\0824\shuttle-v2-core.js：共用 ETA 與資料結構輔助邏輯。"]: bullet(doc,x)
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
